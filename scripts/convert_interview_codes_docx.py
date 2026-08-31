"""
One-off converter: the reconciled-consensus interview-codes Word document
-> the ground-truth CSV template used by core/analytics/validation/
ground_truth_validation.py.

This is a standalone utility script, NOT part of the deployed app -- the
source document's grid-table structure (blank "Codes" cells mean "same
code as the row above"; a category header row sometimes also carries the
first quote for that category, sometimes doesn't) is specific to this one
file, not a general format other documents would follow. Review the
output CSV before treating it as authoritative -- this is a best-effort
automated first pass on a document that was written for a human reader,
not for machine parsing, per the project's own "quick overview, not
research-grade IRR" framing for this data (see PROJECT_STATUS.md).

Usage:
    python scripts/convert_interview_codes_docx.py <input.docx> <output.csv>

Requires: python-docx (already a project dependency).
"""

from __future__ import annotations
import csv
import re
import sys
from pathlib import Path


def convert(docx_path: str, csv_path: str) -> None:
    from docx import Document

    doc = Document(docx_path)
    if not doc.tables:
        raise ValueError(f"No tables found in {docx_path}")
    table = doc.tables[0]

    rows_out = []
    current_category = ""
    current_code = ""
    current_notes = ""
    n_category_rows = 0

    for row in table.rows[1:]:  # skip header row (Codes | Participant quotes | Gr | Researcher Comments)
        cells = row.cells
        codes_cell = cells[0].text.strip() if len(cells) > 0 else ""
        quote_cell = cells[1].text.strip() if len(cells) > 1 else ""
        notes_cell = cells[3].text.strip() if len(cells) > 3 else ""

        if codes_cell:
            cat_match = re.search(
                r"^(.*?)\s*\(category\)\s*$", codes_cell,
                re.IGNORECASE | re.DOTALL,
            )
            if cat_match:
                n_category_rows += 1
                current_category = re.sub(r"\s+", " ", cat_match.group(1)).strip()
                # This source document sometimes uses the category header
                # row to also introduce the first code's quote (no separate
                # code label follows), and sometimes leaves it as a pure
                # section divider with the real code on the next row. Using
                # the category text itself as a fallback code label for
                # this row is the safer default -- worth spot-checking rows
                # where "code_name" ends up equal to "category" in the
                # output, since that's this fallback firing.
                current_code = current_category
            else:
                current_code = re.sub(r"\s+", " ", codes_cell).strip()
            current_notes = re.sub(r"\s+", " ", notes_cell).strip()
        elif notes_cell:
            # Blank Codes cell but new notes text -- a continuation comment
            # for the current code block (observed in the source document,
            # e.g. "#hashtag"-style follow-up notes on later rows).
            extra = re.sub(r"\s+", " ", notes_cell).strip()
            current_notes = f"{current_notes} {extra}".strip() if current_notes else extra

        if not quote_cell:
            continue  # spacer row or category-only row with no quote

        page_match = re.search(r"\(p\s*(\d+)\)", quote_cell, re.IGNORECASE)
        page_ref = page_match.group(1) if page_match else ""

        rows_out.append({
            "category":            current_category,
            "code_name":           current_code,
            "quote":               re.sub(r"\s+", " ", quote_cell).strip(),
            "participant_id":      "",   # deliberately blank -- see module docstring
            "page_ref":            page_ref,
            "researcher_initials": "",   # not present in the source document
            "notes":               current_notes,
        })

    fieldnames = [
        "category", "code_name", "quote",
        "participant_id", "page_ref", "researcher_initials", "notes",
    ]
    with open(csv_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows_out)

    n_no_page_ref = sum(1 for r in rows_out if not r["page_ref"])
    print(f"Wrote {len(rows_out)} rows to {csv_path}")
    print(f"  {n_category_rows} category headers encountered")
    print(f"  {n_no_page_ref} row(s) with no detected page reference (pN)")
    print(
        "  This is a best-effort automated parse -- spot-check the output, "
        "especially rows where code_name matches category exactly (the "
        "category-row-doubles-as-code-name fallback described in this "
        "script's docstring)."
    )


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print(__doc__)
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2])
